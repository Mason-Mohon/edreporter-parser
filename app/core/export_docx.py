"""DOCX and Markdown export functionality.

This module handles exporting extracted articles to DOCX and Markdown formats.
"""

from pathlib import Path
from typing import Dict

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.segment_model import AnnotationDoc
from app.core.extraction import ExtractedText


def export_issue_docx(
    output_path: str | Path,
    annotation_doc: AnnotationDoc,
    article_texts: Dict[str, ExtractedText]
) -> None:
    """Export all articles to a single DOCX file.
    
    Args:
        output_path: Path to output DOCX file
        annotation_doc: Annotation document with article metadata
        article_texts: Extracted text for each article
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add title page info
    doc.add_heading('EdReporter Articles', level=0)
    doc.add_paragraph(f"Source: {Path(annotation_doc.source_pdf).name}")
    doc.add_paragraph(f"Extracted: {annotation_doc.updated_at}")
    doc.add_page_break()
    
    # Sort articles by ID
    article_ids = sorted(annotation_doc.articles.keys())
    
    # Export each article
    for article_id in article_ids:
        article = annotation_doc.articles.get(article_id)
        extracted = article_texts.get(article_id)
        
        if not article:
            continue
        
        # Article title (Heading 1)
        if article.title:
            doc.add_heading(article.title, level=1)
        else:
            doc.add_heading(f"Article {article_id}", level=1)
        
        # Subtitle (Heading 2)
        if article.subtitle:
            doc.add_heading(article.subtitle, level=2)
        
        # Author (italic)
        if article.author:
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(f"By {article.author}")
            author_run.italic = True
            author_para.space_after = Pt(12)
        
        # Tags
        if article.tags:
            tags_para = doc.add_paragraph()
            tags_run = tags_para.add_run(f"Tags: {', '.join(article.tags)}")
            tags_run.font.size = Pt(9)
            tags_para.space_after = Pt(12)
        
        # Body text
        if extracted and extracted.text:
            # Split into paragraphs (separated by blank lines)
            paragraphs = extracted.text.split('\n\n')
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)
        else:
            doc.add_paragraph("[No text extracted]")
        
        # Add separator between articles
        doc.add_paragraph()
        doc.add_paragraph("—" * 40)
        doc.add_paragraph()
    
    # Add appendix with extraction metadata
    doc.add_page_break()
    doc.add_heading('Extraction Metadata', level=1)
    
    doc.add_paragraph(f"Source PDF: {annotation_doc.source_pdf}")
    doc.add_paragraph(f"DPI: {annotation_doc.settings.dpi}")
    doc.add_paragraph(f"OCR Language: {annotation_doc.settings.ocr_lang}")
    doc.add_paragraph(f"Prefer PDF Text: {annotation_doc.settings.prefer_pdf_text_layer}")
    doc.add_paragraph()
    
    # Per-article extraction details
    for article_id in article_ids:
        extracted = article_texts.get(article_id)
        if not extracted:
            continue
        
        article = annotation_doc.articles.get(article_id)
        article_title = article.title if article and article.title else f"Article {article_id}"
        
        doc.add_heading(f"{article_id}: {article_title}", level=2)
        
        # Create table for region details
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Page'
        header_cells[1].text = 'Region ID'
        header_cells[2].text = 'Method'
        header_cells[3].text = 'Quality'
        header_cells[4].text = 'Length'
        
        # Data rows
        for region_meta in extracted.regions_metadata:
            row_cells = table.add_row().cells
            row_cells[0].text = str(region_meta.page_index + 1)  # 1-indexed for display
            row_cells[1].text = region_meta.region_id[:8]  # Shorten UUID
            row_cells[2].text = region_meta.method
            row_cells[3].text = f"{region_meta.quality_score:.2f}"
            row_cells[4].text = str(region_meta.text_length)
        
        doc.add_paragraph()
    
    # Save document
    doc.save(output_path)


def export_markdown(
    output_path: str | Path,
    annotation_doc: AnnotationDoc,
    article_texts: Dict[str, ExtractedText]
) -> None:
    """Export all articles to a Markdown file.
    
    Args:
        output_path: Path to output Markdown file
        annotation_doc: Annotation document with article metadata
        article_texts: Extracted text for each article
    """
    lines = []
    
    # Header
    lines.append("# EdReporter Articles\n")
    lines.append(f"**Source:** {Path(annotation_doc.source_pdf).name}\n")
    lines.append(f"**Extracted:** {annotation_doc.updated_at}\n")
    lines.append("\n---\n")
    
    # Sort articles by ID
    article_ids = sorted(annotation_doc.articles.keys())
    
    # Export each article
    for article_id in article_ids:
        article = annotation_doc.articles.get(article_id)
        extracted = article_texts.get(article_id)
        
        if not article:
            continue
        
        # Article header
        lines.append(f"\n## {article.title or f'Article {article_id}'}\n")
        
        if article.subtitle:
            lines.append(f"### {article.subtitle}\n")
        
        if article.author:
            lines.append(f"*By {article.author}*\n")
        
        if article.tags:
            lines.append(f"**Tags:** {', '.join(article.tags)}\n")
        
        lines.append("\n")
        
        # Body text
        if extracted and extracted.text:
            lines.append(extracted.text)
            lines.append("\n")
        else:
            lines.append("*[No text extracted]*\n")
        
        lines.append("\n---\n")
    
    # Write to file
    output_path = Path(output_path)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))


def export_article_docx(
    output_path: str | Path,
    article_id: str,
    annotation_doc: AnnotationDoc,
    extracted_text: ExtractedText
) -> None:
    """Export a single article to its own DOCX file.
    
    Args:
        output_path: Path to output DOCX file
        article_id: ID of the article to export
        annotation_doc: Annotation document
        extracted_text: Extracted text for the article
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    article = annotation_doc.articles.get(article_id)
    
    if not article:
        return
    
    # Article title
    if article.title:
        doc.add_heading(article.title, level=1)
    else:
        doc.add_heading(f"Article {article_id}", level=1)
    
    # Subtitle
    if article.subtitle:
        doc.add_heading(article.subtitle, level=2)
    
    # Author
    if article.author:
        author_para = doc.add_paragraph()
        author_run = author_para.add_run(f"By {article.author}")
        author_run.italic = True
        author_para.space_after = Pt(12)
    
    # Tags
    if article.tags:
        tags_para = doc.add_paragraph()
        tags_run = tags_para.add_run(f"Tags: {', '.join(article.tags)}")
        tags_run.font.size = Pt(9)
        tags_para.space_after = Pt(12)
    
    # Body text
    if extracted_text and extracted_text.text:
        paragraphs = extracted_text.text.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)
    else:
        doc.add_paragraph("[No text extracted]")
    
    # Add metadata section
    doc.add_page_break()
    doc.add_heading('Metadata', level=2)
    
    doc.add_paragraph(f"Source PDF: {annotation_doc.source_pdf}")
    doc.add_paragraph(f"Article ID: {article_id}")
    
    # Save document
    doc.save(output_path)

